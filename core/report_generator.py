from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_vuln_table(document, vuln):
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(title, value):
        row = table.add_row().cells
        row[0].text = title
        row[1].text = value

    add_row("Targeted IP", vuln.get("targeted_ip", "N/A"))
    add_row("Source IP", vuln.get("source_ip", "N/A"))
    add_row("Method Used", vuln.get("method_used", "N/A"))
    add_row("Port Used", vuln.get("port_used", "N/A"))
    add_row("Vulnerability Details", vuln.get("vulnerability_details", "N/A"))
    add_row("Attack Vector", vuln.get("attack_vector", "N/A"))
    add_row("Attack Complexity", vuln.get("attack_complexity", "N/A"))
    add_row("Privileges Required", vuln.get("privileges_required", "N/A"))
    add_row("User Interaction", vuln.get("user_interaction", "N/A"))
    add_row("Scope", vuln.get("scope", "N/A"))
    add_row("Confidentiality Impact", vuln.get("confidentiality", "N/A"))
    add_row("Integrity Impact", vuln.get("integrity", "N/A"))
    add_row("Availability Impact", vuln.get("availability", "N/A"))
    add_row("Severity Rating", vuln.get("severity_rating", "N/A"))
    add_row("Business Impact", vuln.get("business_impact", "N/A"))
    add_row("Remediation", vuln.get("remediation", "N/A"))
    add_row("Proof of Concept", vuln.get("proof_of_concept", "N/A"))

    document.add_paragraph("")  # add spacing

def report_generator(data, output_path, fmt="docx", tables=None):
    document = Document()
    document.add_heading("Security Assessment Report", level=0)

    if isinstance(data, dict):
        data = [data]

    for vuln in data:
        format_heading(document, vuln.get("issue_name", "Vulnerability"), level=2)
        add_vuln_table(document, vuln)

    if tables:
        for table_info in tables:
            document.add_heading(table_info["title"], level=2)
            rows = table_info["data"]
            table = document.add_table(rows=1, cols=len(rows[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(rows[0]):
                hdr_cells[i].text = h
            for row in rows[1:]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = val

    document.save(output_path)
    print(f"[+] Report saved to {output_path}")
